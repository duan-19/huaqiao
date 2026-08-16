from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(r"C:\Users\29969\Documents\Codex\2026-08-11\ui-ux-ai-1-ppt-gis")
SOURCE = ROOT / "work" / "docx-edit" / "source-template.docx"
OUTPUT = ROOT / "outputs" / "“青声传薪，数扬侨魂”华侨革命史数字科普实践平台项目申报书.docx"


def set_run_font(run, font_name="仿宋_GB2312", size=14, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), font_name)


def set_para_layout(paragraph, line=20, indent=False, keep_next=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(line)
    fmt.first_line_indent = Pt(28) if indent else Pt(0)
    ppr = paragraph._p.get_or_add_pPr()
    keep = ppr.find(qn("w:keepNext"))
    if keep_next and keep is None:
        ppr.append(OxmlElement("w:keepNext"))
    elif not keep_next and keep is not None:
        ppr.remove(keep)


def clear_cell_paragraphs(cell):
    for child in list(cell._tc):
        if child.tag == qn("w:p"):
            cell._tc.remove(child)


def add_para(cell, text, kind="body"):
    paragraph = cell.add_paragraph()
    if kind == "major":
        set_para_layout(paragraph, line=22, indent=False, keep_next=True)
        run = paragraph.add_run(text)
        set_run_font(run, "黑体", 14)
    elif kind == "sub":
        set_para_layout(paragraph, line=20, indent=False, keep_next=True)
        run = paragraph.add_run(text)
        set_run_font(run, "仿宋_GB2312", 14)
    elif kind == "stage":
        set_para_layout(paragraph, line=20, indent=False)
        run = paragraph.add_run(text)
        set_run_font(run, "仿宋_GB2312", 14)
    else:
        set_para_layout(paragraph, line=20, indent=True)
        run = paragraph.add_run(text)
        set_run_font(run, "仿宋_GB2312", 14)
    return paragraph


def replace_cell_content(cell, paragraphs):
    clear_cell_paragraphs(cell)
    for text, kind in paragraphs:
        add_para(cell, text, kind)


def update_cover(document):
    paragraph = document.paragraphs[8]
    label = "项目名称："
    value = "“青声传薪，数扬侨魂”华侨革命史数字科普实践平台"

    # 保留原模板的段落制表位、行距和下划线布局，仅重建文字运行。
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        label_run = paragraph.runs[0]
    else:
        label_run = paragraph.add_run()
    label_run.text = label
    set_run_font(label_run, "黑体", 16)

    value_run = paragraph.add_run(value)
    # 全称较长，使用与模板一致的仿宋字体并适度缩小，确保封面单行排版，
    # 避免将模板中的分页符挤到单独空白页。
    set_run_font(value_run, "仿宋_GB2312", 12)
    value_run.font.underline = True


BACKGROUND = [
    ("一、实践项目背景及目标", "major"),
    ("（一）项目概述", "sub"),
    (
        "本项目面向社区青少年、学校教师、社区工作者及关注华侨历史的社会公众，依托团队真实社会实践建设公益性华侨革命史数字科普平台。项目不是一次性的成果照片展示，而是将三期社区主题宣讲、场馆寻访、史料学习、实践影像、原创课件和实践报道转化为可持续使用的线上学习资源。平台围绕“社区传薪、寻访侨迹、溯源侨史、数字留存、长效传播”的主线，设置三堂侨课、寻访侨迹、数字侨史、实践成果、数字资源和知识答题等模块，支持课件在线浏览与公益下载、照片画廊、互动时间轴、权威学习入口等功能。",
        "body",
    ),
    ("（二）项目背景与需求", "sub"),
    (
        "华侨革命史承载着跨越山海的爱国记忆，是开展青少年理想信念教育和家国情怀教育的重要资源。当前相关内容多分散于展馆、文献和专业研究中，面向青少年的表达较少；部分社会实践成果也停留在单次宣讲、照片汇总或短期报道层面，活动结束后难以继续服务学校和社区。团队希望发挥计算机专业优势，用清晰、轻量、可核实的数字化方式连接历史档案与当代青年，让青少年能够看懂、愿意阅读，并能继续获取和使用相关资源。",
        "body",
    ),
    ("（三）项目目标", "sub"),
    (
        "项目拟实现三个目标：一是以三期社区宣讲为内容入口，用问题式、故事化、适龄化表达讲清抗战华侨力量、青年华侨先烈和侨批家国记忆；二是在宣讲后走进华侨博物院及侨批相关场馆，通过真实展品、文献和现场学习进一步核实、丰富宣讲内容，所有历史信息尽可能标明来源，未核实信息不作补全；三是建设便于维护、适配电脑与手机的静态科普网站，形成可长期浏览、传播和更新的公益数字资源，体现“计算机专业、社会实践、华侨革命史传播”的融合特色。",
        "body",
    ),
]


PLAN = [
    ("二、实践方案", "major"),
    ("（一）实践地点基本情况", "sub"),
    (
        "本项目以福建省厦门市思明区、集美区为主要实践区域。社区宣讲依托侨英街道凤林美社区开展，面向青少年进行华侨革命史主题科普；三期宣讲完成后，团队再前往华侨博物院和侨批相关场馆参观学习、记录史料线索。社区提供面向青少年的真实传播场景，场馆提供实物、文献和专业讲解，两类场景共同构成由课堂问题走向历史现场的实践闭环。",
        "body",
    ),
    ("（二）实践实施路径", "sub"),
    (
        "第一阶段：社区传薪。团队完成内容准备和分工后，依次开展《烽火侨心——抗战华侨力量》《少年榜样·青年华侨先烈》《寻访侨批记忆 厚植少年家国情怀》三期社区青少年主题宣讲。通过PPT讲解、问题引导和现场互动，帮助青少年理解抗战华侨力量、青年担当以及侨批中的乡情与家国情怀。",
        "stage",
    ),
    (
        "第二阶段：寻访侨迹。三期社区宣讲完成后，团队带着课堂中形成的问题走进华侨博物院和侨批相关场馆，通过参观展览、聆听讲解、观察实物、记录资料来源等方式，对宣讲内容进行进一步学习、印证和补充。",
        "stage",
    ),
    (
        "第三阶段：数字留存。团队整合三期宣讲PPT、社区活动照片、场馆参访照片、实践推文及经核实的相关资料，完成内容分类、图片压缩、来源标注和网页制作，建设能够长期开放的华侨革命史数字科普实践平台。",
        "stage",
    ),
    ("（三）主要实践内容与方式", "sub"),
    (
        "1.开展三堂青少年侨史课。第一课以“远在海外的华侨为什么义无反顾支援祖国抗战”为问题入口，介绍捐资支援、海外宣传、归国服务等抗战华侨力量；第二课聚焦与今天青年年龄相仿的华侨先烈，引导青少年思考理想、责任与担当；第三课从跨越山海的侨批出发，理解海外游子的乡情、亲情和家国情怀。宣讲不照搬史料或课件，而是以清晰叙事、图像材料和互动问答降低理解门槛。",
        "body",
    ),
    (
        "2.开展场馆参访和史料寻访。团队在完成社区宣讲后进入历史现场，重点关注与三堂侨课相关的展品、文献、侨批和人物线索，并记录来源机构、资料名称、参考链接和访问情况。对暂未获得馆藏编号或来源不明确的内容，不进行推测和编造，在平台中统一标记为“史料来源待核实”。",
        "body",
    ),
    (
        "3.建设公益数字科普平台。网站采用前端静态技术建设，不设置注册、登录、评论、支付和数据库，降低后期维护成本。首页用三十秒左右的浏览路径讲清团队做了什么；三期侨课页面集中展示课程内容、活动照片和PPT/PDF；寻访侨迹页面呈现场馆学习过程；数字侨史页面通过时间轴和史料卡片开展轻量科普；成果与资源页面提供实践推文、照片、视频、官方展馆和权威学习资源入口；知识答题和前端搜索提升青少年的参与感与信息获取效率。",
        "body",
    ),
    (
        "4.开展数字资料整理。团队对图片进行分类、压缩和替换路径管理，首屏图片与相册缩略图分别优化；对三套宣讲课件同步保留源文件和PDF版本，便于在线浏览与公益下载；课程、场馆、文章、时间轴、资源和题库内容采用独立数据文件维护，使后续成员无需修改页面程序即可更新内容。",
        "body",
    ),
    (
        "5.推进成果传播与服务。平台优先适配手机和电脑浏览，可用于社区活动后的延伸学习、学校主题班会或团课、实践成果展示及青少年自主阅读。团队将结合真实实践推文、校院和社区传播渠道，引导访问者从“看见项目”逐步进入“了解、阅读、互动、获取资源”的完整学习路径。",
        "body",
    ),
    ("（四）项目特色与差异化", "sub"),
    (
        "项目坚持真实实践逻辑，先在社区完成三期宣讲，再走进场馆寻找答案，最后将线下实践整理为数字资源，避免将流程包装为与事实不符的“先调研后宣讲”。与普通照片墙式成果页面相比，本项目突出青年身份和青少年表达，以“旧侨批与新青年、历史档案与数字网页、百年前青年与今天青年、远隔山海与家国同心”为核心视觉和叙事关系；同时设置来源组件、资料核实状态和可维护的数据结构，在视觉美观、内容严谨、公益使用和长期更新之间形成平衡。",
        "body",
    ),
    ("（五）人员分工", "sub"),
    (
        "段佳锐负责总体统筹、进度协调、经费管理、对接社区与指导教师、成果汇总及网页建设协调；廖彦负责史料整理、三期宣讲课件和网页科普内容；贾玲玲负责主题宣讲、课件制作和青少年化表达；杨佩俊负责活动影像拍摄、素材整理和网页图片呈现；方佳怡负责三期活动主持、主题宣讲和现场互动衔接；张婉婷负责主题宣讲、视频剪辑和实践内容整理；熊媛媛负责活动场地与物资协调、场馆资料和照片档案整理。全体成员共同参与网站内容核验、测试和后续维护。",
        "body",
    ),
    ("（六）风险与应急预案", "sub"),
    (
        "1.天气、交通与健康风险。提前关注天气和交通信息，明确集合、签到、联络和应急安排；如遇极端天气或成员身体不适，及时调整参访或活动时间并向社区、学院和指导教师报备。",
        "body",
    ),
    (
        "2.活动现场与青少年安全。遵守社区和场馆秩序，不擅自触碰展品；由团队负责人统一组织人员与设备，涉及青少年影像展示时遵循授权和隐私保护要求。",
        "body",
    ),
    (
        "3.设备与资料风险。宣讲课件、网站代码、照片和视频采用本地与云端双重备份，关键课件另存PDF；现场设备提前测试，网络不可用时使用离线课件，避免影响宣讲。",
        "body",
    ),
    (
        "4.内容与版权风险。历史内容须标明来源，优先使用官方展馆、权威出版物或已获授权资料；不编造史实、馆藏编号、人物经历和参与者评价。外部资源统一标注来源并以链接方式提供，网站同步发布公益科普与版权说明。",
        "body",
    ),
    (
        "5.网站运行风险。首版采用轻量静态架构和响应式设计，减少复杂依赖；上线前进行多浏览器和手机端测试，压缩图片并设置替代文字。正式公开访问时配置稳定托管、独立域名、HTTPS、基础搜索引擎信息和定期备份机制。",
        "body",
    ),
]


RESULTS = [
    ("三、预期成果与成效", "major"),
    ("（一）已有实践基础", "sub"),
    (
        "团队已完成三期社区青少年主题宣讲，并在此基础上开展华侨博物院和侨批相关场馆的线下参观、学习与史料寻访；现已进入三套宣讲课件、活动照片、场馆影像、实践推文和相关资料的系统整理阶段，同时推进数字科普网页首版建设。上述真实实践为平台内容提供了稳定来源，也形成了由社区课堂走向历史现场、再回到数字传播的完整叙事基础。",
        "body",
    ),
    ("（二）预期成果", "sub"),
    (
        "1.形成三套原创主题宣讲课件及对应PDF版本，分别围绕抗战华侨力量、青年华侨先烈和侨批家国记忆，支持在线浏览、公益下载和社区、学校教育场景复用。",
        "body",
    ),
    (
        "2.建成“青声传薪，数扬侨魂”华侨革命史数字科普实践平台，完成首页、三期侨课、寻访侨迹、数字侨史、实践成果、数字资源、知识答题和前端搜索等核心模块，实现电脑与手机端响应式访问。",
        "body",
    ),
    (
        "3.建立实践照片与场馆影像档案，整理真实实践推文、短视频及活动记录，形成可持续补充的数字成果目录；所有图片设置说明和替代文字，历史类内容统一标注资料来源或核实状态。",
        "body",
    ),
    (
        "4.建设轻量科普内容，包括华侨革命史互动时间轴、人物与侨批史料卡片、场馆官方学习入口、权威资源链接和趣味知识题库，使访问者不仅能够观看成果，还能继续阅读、互动和获取资源。",
        "body",
    ),
    ("（三）服务对象与应用场景", "sub"),
    (
        "项目主要服务社区青少年，同时面向学校教师、社区工作者、大学生社会实践团队及关注华侨历史的社会公众。平台可用于主题班会、团课、社区公益课堂、校内实践成果展示和自主学习；教师或社区工作者能够直接浏览或下载课件，普通访问者能够通过手机查看故事、照片、推文和权威学习资源。",
        "body",
    ),
    ("（四）长效机制与社会成效", "sub"),
    (
        "网站以数据与页面分离的方式建设，课程、场馆、文章、史料、资源和题库均可独立更新，降低成员更替后的维护门槛。团队将持续核实资料来源、补充真实实践成果、维护外部链接和优化移动端体验，使一次暑期社会实践沉淀为长期开放的公益科普资源。项目预期帮助青少年从具体人物、侨批和历史选择中理解华侨爱国精神，也展示计算机专业学生以数字技术服务文化传播和社会教育的实践能力。",
        "body",
    ),
]


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SOURCE, OUTPUT)
    document = Document(OUTPUT)

    update_cover(document)
    main_table = document.tables[0]
    replace_cell_content(main_table.rows[14].cells[0], BACKGROUND)
    replace_cell_content(main_table.rows[15].cells[0], PLAN)
    replace_cell_content(main_table.rows[16].cells[0], RESULTS)

    # 延续模板的“与下段同页”设置，避免标题孤立在页尾。
    for cell in (main_table.rows[14].cells[0], main_table.rows[15].cells[0], main_table.rows[16].cells[0]):
        for paragraph in cell.paragraphs:
            if paragraph.text.startswith(("一、", "二、", "三、", "（")):
                ppr = paragraph._p.get_or_add_pPr()
                if ppr.find(qn("w:keepNext")) is None:
                    ppr.append(OxmlElement("w:keepNext"))

    document.core_properties.title = "“青声传薪，数扬侨魂”华侨革命史数字科普实践平台项目申报书"
    document.core_properties.subject = "2026年暑期社会实践活动团队项目申报书"
    document.core_properties.keywords = "华侨革命史,青少年科普,社会实践,数字平台"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
