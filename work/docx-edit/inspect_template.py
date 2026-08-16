from docx import Document
from docx.oxml.ns import qn
from pathlib import Path

doc_path = Path('work/docx-edit/source-template.docx')
doc = Document(doc_path)

print('PARAGRAPHS', len(doc.paragraphs), 'TABLES', len(doc.tables), 'SECTIONS', len(doc.sections))
for i, p in enumerate(doc.paragraphs):
    text = p.text.replace('\n', ' / ').strip()
    if text:
        print(f'P{i:03d} [{p.style.name}] {text}')

for ti, table in enumerate(doc.tables):
    print(f'\n=== TABLE {ti} rows={len(table.rows)} cols={len(table.columns)} ===')
    for ri, row in enumerate(table.rows):
        values = []
        for ci, cell in enumerate(row.cells):
            text = ' / '.join(p.text.replace('\n', ' ').strip() for p in cell.paragraphs if p.text.strip())
            values.append(f'C{ci}:{text}')
        print(f'R{ri:02d} | ' + ' || '.join(values))

print('\nSTYLES USED')
seen = {}
for p in doc.paragraphs:
    seen[p.style.name] = seen.get(p.style.name, 0) + 1
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                seen[p.style.name] = seen.get(p.style.name, 0) + 1
for name, count in sorted(seen.items(), key=lambda x: (-x[1], x[0])):
    print(name, count)
