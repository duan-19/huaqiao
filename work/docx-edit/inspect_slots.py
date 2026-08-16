from docx import Document
from docx.oxml.ns import qn
from pathlib import Path

doc = Document('work/docx-edit/source-template.docx')
table = doc.tables[0]

def fmt_value(run):
    rpr = run._element.rPr
    east = rpr.rFonts.get(qn('w:eastAsia')) if rpr is not None and rpr.rFonts is not None else None
    ascii_font = rpr.rFonts.get(qn('w:ascii')) if rpr is not None and rpr.rFonts is not None else None
    return {
        'text': run.text,
        'font': run.font.name,
        'east': east,
        'ascii': ascii_font,
        'size': run.font.size.pt if run.font.size else None,
        'bold': run.bold,
        'italic': run.italic,
    }

for row_index in [14, 15, 16, 17]:
    cell = table.cell(row_index, 0)
    print(f'\nROW {row_index} unique_tc={id(cell._tc)} paragraphs={len(cell.paragraphs)}')
    for pi, p in enumerate(cell.paragraphs):
        pf = p.paragraph_format
        print('P', pi, 'style', p.style.name, 'align', p.alignment, 'left', pf.left_indent.pt if pf.left_indent else None,
              'first', pf.first_line_indent.pt if pf.first_line_indent else None, 'before', pf.space_before.pt if pf.space_before else None,
              'after', pf.space_after.pt if pf.space_after else None, 'line', pf.line_spacing)
        print(' text=', repr(p.text))
        for ri, run in enumerate(p.runs):
            if run.text:
                print('  R', ri, fmt_value(run))

print('\nCOVER')
for idx in [5,6,8,9,10,11,16,18]:
    p = doc.paragraphs[idx]
    print('P', idx, repr(p.text), 'align', p.alignment, 'style', p.style.name)
    for ri, run in enumerate(p.runs):
        if run.text:
            print('  R',ri,fmt_value(run))
