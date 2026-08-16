from io import BytesIO
from pathlib import Path
import zipfile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\29969\Documents\Codex\2026-08-11\ui-ux-ai-1-ppt-gis")
SOURCE = ROOT / "work" / "docx-edit" / "source-template.docx"
OUTPUT = ROOT / "outputs" / "青声传薪数扬侨魂华侨革命史数字科普实践平台项目申报书（目录版）.docx"

FONT_BODY = "仿宋_GB2312"
FONT_HEADING = "黑体"
FONT_SONG = "宋体"
FONT_COVER = "方正小标宋简体"
COLOR_RED = RGBColor(139, 46, 46)
COLOR_DARK = RGBColor(39, 37, 34)
COLOR_GRAY = RGBColor(105, 100, 93)


def set_east_asia_font(run_or_style, name):
    run_or_style.font.name = name
    rpr = run_or_style._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)


def set_run(run, font=FONT_BODY, size=12, bold=False, color=COLOR_DARK):
    set_east_asia_font(run, font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    return run


def set_style_font(style, font, size, bold=False, color=COLOR_DARK):
    set_east_asia_font(style, font)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = color


def remove_existing_body(document):
    body = document._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def configure_page(document):
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)
    section.header_distance = Cm(1.3)
    section.footer_distance = Cm(1.4)


def configure_styles(document):
    styles = document.styles

    def get_or_add_paragraph_style(name):
        if name in styles:
            return styles[name]
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]
        style.next_paragraph_style = styles["Normal"]
        return style

    def set_outline_level(style, level):
        ppr = style._element.get_or_add_pPr()
        old = ppr.find(qn("w:outlineLvl"))
        if old is not None:
            ppr.remove(old)
        outline = OxmlElement("w:outlineLvl")
        outline.set(qn("w:val"), str(level))
        ppr.append(outline)
        if style._element.find(qn("w:qFormat")) is None:
            style._element.append(OxmlElement("w:qFormat"))

    normal = styles["Normal"]
    set_style_font(normal, FONT_BODY, 12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(22)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.widow_control = True

    heading1 = get_or_add_paragraph_style("Heading 1")
    set_outline_level(heading1, 0)
    set_style_font(heading1, FONT_HEADING, 16, bold=True)
    heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading1.paragraph_format.first_line_indent = Pt(0)
    heading1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    heading1.paragraph_format.line_spacing = Pt(28)
    heading1.paragraph_format.space_before = Pt(12)
    heading1.paragraph_format.space_after = Pt(6)
    heading1.paragraph_format.keep_with_next = True
    heading1.paragraph_format.keep_together = True

    heading2 = get_or_add_paragraph_style("Heading 2")
    set_outline_level(heading2, 1)
    set_style_font(heading2, FONT_HEADING, 14, bold=True)
    heading2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading2.paragraph_format.first_line_indent = Pt(0)
    heading2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    heading2.paragraph_format.line_spacing = Pt(25)
    heading2.paragraph_format.space_before = Pt(8)
    heading2.paragraph_format.space_after = Pt(4)
    heading2.paragraph_format.keep_with_next = True
    heading2.paragraph_format.keep_together = True

    heading3 = get_or_add_paragraph_style("Heading 3")
    set_outline_level(heading3, 2)
    set_style_font(heading3, FONT_HEADING, 12, bold=True)
    heading3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading3.paragraph_format.left_indent = Pt(12)
    heading3.paragraph_format.first_line_indent = Pt(0)
    heading3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    heading3.paragraph_format.line_spacing = Pt(23)
    heading3.paragraph_format.space_before = Pt(5)
    heading3.paragraph_format.space_after = Pt(2)
    heading3.paragraph_format.keep_with_next = True
    heading3.paragraph_format.keep_together = True

    for name, size, left, line in (
        ("TOC 1", 10.5, 0, 16),
        ("TOC 2", 10, 18, 15),
        ("TOC 3", 9.5, 36, 14),
    ):
        style = get_or_add_paragraph_style(name)
        set_style_font(style, FONT_SONG, size)
        style.paragraph_format.left_indent = Pt(left)
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        style.paragraph_format.line_spacing = Pt(line)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)


def add_bottom_border(paragraph, color="8B2E2E", size="10", space="6"):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)


def add_cover(document, logo_bytes):
    logo_p = document.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_p.paragraph_format.space_before = Pt(24)
    logo_p.paragraph_format.space_after = Pt(38)
    # 图片为行内对象，必须覆盖正文“固定22磅”行距，否则校徽会被压成细线。
    logo_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    logo_p.paragraph_format.line_spacing = 1
    logo_p.paragraph_format.first_line_indent = Pt(0)
    logo_p.add_run().add_picture(BytesIO(logo_bytes), width=Inches(4.15))

    activity = document.add_paragraph()
    activity.alignment = WD_ALIGN_PARAGRAPH.CENTER
    activity.paragraph_format.space_after = Pt(10)
    set_run(activity.add_run("2026年暑期社会实践"), FONT_HEADING, 20, bold=False)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    title.paragraph_format.line_spacing = Pt(45)
    title.paragraph_format.space_after = Pt(30)
    set_run(title.add_run("项 目 申 报 书"), FONT_HEADING, 32, bold=True)

    project = document.add_paragraph()
    project.alignment = WD_ALIGN_PARAGRAPH.CENTER
    project.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    project.paragraph_format.line_spacing = Pt(30)
    project.paragraph_format.space_after = Pt(8)
    set_run(project.add_run("“青声传薪，数扬侨魂”\n华侨革命史数字科普实践平台"), FONT_COVER, 20, bold=False, color=COLOR_RED)
    add_bottom_border(project)

    slogan = document.add_paragraph()
    slogan.alignment = WD_ALIGN_PARAGRAPH.CENTER
    slogan.paragraph_format.space_before = Pt(12)
    slogan.paragraph_format.space_after = Pt(36)
    set_run(slogan.add_run("社区传薪 · 寻访侨迹 · 溯源侨史 · 数字留存 · 长效传播"), FONT_SONG, 10.5, color=COLOR_GRAY)

    metadata = [
        ("申报团队", "“青声传薪，数扬侨魂”华侨革命史科普实践团"),
        ("所属单位", "华侨大学计算机科学与技术学院"),
        ("项目负责人", "段佳锐"),
        ("指导教师", "黄毅凯"),
    ]
    for label, value in metadata:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(28)
        set_run(p.add_run(f"{label}："), FONT_HEADING, 14)
        set_run(p.add_run(value), FONT_BODY, 14)

    date_p = document.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.space_before = Pt(38)
    set_run(date_p.add_run("2026年8月"), FONT_SONG, 12)
    date_p.add_run().add_break(WD_BREAK.PAGE)


def add_toc_field(paragraph):
    paragraph.paragraph_format.first_line_indent = Pt(0)
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    instr_run._r.append(instr)

    sep_run = paragraph.add_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    sep_run._r.append(sep)

    display = paragraph.add_run("目录将在打开文档时自动更新")
    set_run(display, FONT_SONG, 10, color=COLOR_GRAY)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def add_toc(document):
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Pt(0)
    title.paragraph_format.space_before = Pt(4)
    title.paragraph_format.space_after = Pt(8)
    set_run(title.add_run("目  录"), FONT_HEADING, 18, bold=True)

    toc = document.add_paragraph()
    add_toc_field(toc)

    spacer = document.add_paragraph()
    spacer.paragraph_format.first_line_indent = Pt(0)
    spacer.add_run().add_break(WD_BREAK.PAGE)


def add_heading(document, text, level):
    return document.add_paragraph(text, style=f"Heading {level}")


def add_body(document, text):
    p = document.add_paragraph(style="Normal")
    p.add_run(text)
    return p


def add_labeled_body(document, label, text):
    p = document.add_paragraph(style="Normal")
    set_run(p.add_run(label), FONT_BODY, 12, bold=True)
    set_run(p.add_run(text), FONT_BODY, 12)
    return p


def add_content(document):
    add_heading(document, "第一章  项目概述", 1)
    add_body(
        document,
        "“青声传薪，数扬侨魂”华侨革命史数字科普实践平台，是华侨大学计算机科学与技术学院学生团队围绕华侨革命史开展的青年公益科普项目。项目以社区青少年为主要服务对象，用青年能够讲清、少年能够听懂的方式呈现抗战华侨力量、青年华侨先烈和侨批中的家国记忆。",
    )
    add_body(
        document,
        "项目实践遵循真实发生顺序：首先在侨英街道凤林美社区连续开展三期青少年主题宣讲；随后带着课堂中的问题走进华侨博物院和侨批相关场馆，学习展陈内容并记录史料线索；最后整合三期宣讲课件、活动照片、场馆影像、实践推文和相关资料，建设长期开放的数字科普网页。",
    )
    add_body(
        document,
        "项目希望将一次暑期社会实践转化为可持续使用的公共文化资源，形成“社区传薪、寻访侨迹、溯源侨史、数字留存、长效传播”的完整路径，体现计算机专业青年以数字技术服务红色文化传播和社会教育的实践特色。",
    )

    add_heading(document, "第二章  项目背景", 1)
    add_heading(document, "2.1  文化与育人背景", 2)
    add_body(
        document,
        "华侨革命史记录了海外侨胞在民族危亡和国家建设进程中的家国选择，是开展爱国主义教育、理想信念教育和青年担当教育的重要内容。厦门具有深厚的侨乡文化积淀，社区、展馆和侨批史料为青少年理解华侨爱国精神提供了真实而具体的学习场景。高校社会实践团队有责任把这些历史资源转化为更贴近当代青少年的表达。",
    )

    add_heading(document, "2.2  现实需求", 2)
    add_body(
        document,
        "目前，华侨革命史资料较多分散于展馆、文献和专业研究中，部分内容专业性较强，青少年在自主接触时存在理解门槛。与此同时，常见社会实践成果多停留在一次宣讲、照片汇总或短期报道层面，活动结束后难以继续为学校和社区提供资源。项目因此需要同时解决“怎样讲得明白”和“怎样长期留下”两个问题。",
    )

    add_heading(document, "2.3  团队优势", 2)
    add_body(
        document,
        "团队成员来自计算机科学与技术学院，具备网页开发、图像处理、视频剪辑和数字内容整理能力，同时已经参与三期社区宣讲和场馆寻访，对项目内容、受众特点和实践过程有直接认识。计算机专业能力与真实线下实践相结合，使团队能够在不脱离史实的前提下完成青少年化表达和数字化转化。",
    )

    add_heading(document, "第三章  项目目标", 1)
    add_heading(document, "3.1  总体目标", 2)
    add_body(
        document,
        "以青年之声讲述侨史，以数字之力传承侨魂。团队通过社区宣讲、场馆寻访和网页建设，把跨越山海的华侨家国故事转化为青少年愿意看、能够懂、可以继续使用的公益科普资源，逐步形成线上线下相衔接的华侨革命史学习平台。",
    )

    add_heading(document, "3.2  具体目标", 2)
    add_labeled_body(document, "一是完成青年化科普表达。", "围绕三个主题形成结构清晰的宣讲内容，用问题、人物、家书和历史选择帮助青少年理解华侨爱国精神。")
    add_labeled_body(document, "二是加强历史现场学习。", "在三期宣讲之后进入华侨博物院和侨批相关场馆，通过展品、文献和现场讲解核实并丰富课程内容。")
    add_labeled_body(document, "三是形成长期数字成果。", "将课件、照片、推文、视频和科普资料整理为适配电脑与手机的静态网页，提供浏览、学习、互动和资源获取入口。")

    add_heading(document, "第四章  项目内容", 1)
    add_heading(document, "4.1  三期社区青少年主题宣讲", 2)
    add_body(
        document,
        "项目从社区课堂出发，在侨英街道凤林美社区依次开展三期主题宣讲。三堂侨课从战争年代、青年榜样和山海家书三个角度展开，形成由国家历史到青年选择、再到普通家庭情感的递进关系。宣讲采用PPT讲解、问题引导和现场互动等方式，避免简单堆叠历史材料。",
    )

    add_heading(document, "4.1.1  烽火侨心——抗战华侨力量", 3)
    add_body(
        document,
        "第一期宣讲围绕“远在海外的华侨为什么会义无反顾支援祖国抗战”展开，介绍海外华侨通过捐资支援、宣传动员、归国服务等方式投身抗战的历史，引导青少年理解跨越地域的民族责任与家国情怀。",
    )

    add_heading(document, "4.1.2  少年榜样·青年华侨先烈", 3)
    add_body(
        document,
        "第二期宣讲把目光投向民族危亡时期挺身而出的华侨青年，通过青年人物和人生选择，拉近百年前青年与今天青少年的距离，引导大家思考理想、责任、勇气与担当。",
    )

    add_heading(document, "4.1.3  寻访侨批记忆 厚植少年家国情怀", 3)
    add_body(
        document,
        "第三期宣讲以一封封跨越山海的侨批为线索，从寄件地、收件地和家书内容出发，讲述海外游子对亲人、故乡和祖国的牵挂，让青少年从具体而温暖的生活记忆中理解“家”与“国”的联系。",
    )

    add_heading(document, "4.2  场馆参访与史料寻访", 2)
    add_body(
        document,
        "三期社区宣讲完成后，团队不把课堂内容作为实践终点，而是带着宣讲中形成的问题走进历史现场。场馆寻访重点关注与三堂侨课相关的展品、人物、侨批和文献线索，通过参观、记录和讨论进一步校正表达、补充知识。",
    )

    add_heading(document, "4.2.1  华侨博物院参访学习", 3)
    add_body(
        document,
        "团队在华侨博物院通过展陈内容和现场讲解，进一步认识海外华侨与祖国命运之间的联系，并将课堂中的文字材料与展柜中的实物、照片和文献相互印证。参访记录将用于完善课程说明和网页科普内容。",
    )

    add_heading(document, "4.2.2  侨批相关场馆寻访", 3)
    add_body(
        document,
        "团队在侨批相关场馆重点了解侨批的形成、流转、内容和历史价值，关注家书背后的家庭故事与社会记忆。涉及史料名称、来源机构、馆藏编号等信息时，坚持以实际记录为准；暂无可靠来源的内容统一标注为“史料来源待核实”，不作推测和编造。",
    )

    add_heading(document, "4.3  数字科普网页建设与成果传播", 2)
    add_body(
        document,
        "线下实践结束后，团队整合三期宣讲PPT、活动照片、场馆参访照片、实践推文及相关资料，建设“青声传薪，数扬侨魂”华侨革命史数字科普实践平台。网页以现代青年视觉与历史档案质感相结合，重点呈现三堂侨课、寻访侨迹、数字侨史、实践成果和数字资源。",
    )
    add_body(
        document,
        "平台采用Vue 3与Vite构建静态网站，不设置注册、登录、评论、支付和数据库，降低后期维护难度。网页支持PPT或PDF在线浏览与公益下载、照片画廊、互动时间轴、知识答题、前端搜索、外部权威学习资源跳转和手机端响应式浏览。所有课程、场馆、文章、史料和资源信息尽量采用独立数据文件维护，方便团队后续更新。",
    )

    add_heading(document, "第五章  实践成果", 1)
    add_heading(document, "5.1  三套原创宣讲课件", 2)
    add_body(
        document,
        "围绕三期社区宣讲，团队形成《烽火侨心——抗战华侨力量》《少年榜样·青年华侨先烈》《寻访侨批记忆 厚植少年家国情怀》三套主题课件。后续将同步整理为PDF版本，便于学校、社区和青少年在线浏览及公益使用。",
    )

    add_heading(document, "5.2  实践影像与报道", 2)
    add_body(
        document,
        "团队对社区活动照片、场馆参访照片、视频素材和实践推文进行分类整理，逐步建立可检索的实践影像档案。影像资料既记录团队实践过程，也作为课程回顾、网页展示和后续传播的真实依据。",
    )

    add_heading(document, "5.3  数字科普平台", 2)
    add_body(
        document,
        "数字网页将线下实践成果汇集到统一入口，使访问者能够在一个平台内看见项目、了解三堂侨课、阅读场馆故事、查看史料来源、参与知识答题并获取学习资源。平台建成后可继续补充真实资料，避免社会实践成果随着活动结束而停止传播。",
    )

    add_heading(document, "第六章  团队介绍", 1)
    add_heading(document, "6.1  团队构成", 2)
    add_body(
        document,
        "实践团由段佳锐、廖彦、贾玲玲、杨佩俊、方佳怡、张婉婷、熊媛媛七名学生组成。团队根据项目需要分工开展总体协调、史料整理、主题宣讲、课件制作、主持互动、影像拍摄、视频剪辑、物资场地协调、网页开发和内容维护，并共同参与历史内容核验与成果整理。",
    )

    add_heading(document, "6.2  指导力量", 2)
    add_body(
        document,
        "项目由计算机科学与技术学院黄毅凯老师指导。指导教师围绕实践方案、活动安全、过程组织、材料规范和成果总结提供指导，帮助团队在真实、规范和可执行的基础上推进项目。",
    )

    add_heading(document, "6.3  团队特色", 2)
    add_body(
        document,
        "团队既是华侨革命史的学习者，也是面向青少年的讲述者和数字内容建设者。成员年龄接近青年受众，能够从今天青年的问题出发重新理解百年前青年的选择；计算机专业背景又使团队具备把课件、影像和史料转化为网页资源的能力，形成内容实践与技术实践相互支撑的团队特色。",
    )

    add_heading(document, "第七章  项目创新点", 1)
    add_heading(document, "7.1  主体创新", 2)
    add_body(
        document,
        "项目突出青年主体身份，由计算机专业大学生面向社区青少年讲述华侨革命史，形成“今天青年讲述百年前青年、青年带动少年”的传播关系，使历史教育更具亲近感和代际连接。",
    )

    add_heading(document, "7.2  内容创新", 2)
    add_body(
        document,
        "项目不追求大而全的历史数据库，而是从三堂侨课出发，以抗战华侨力量、青年华侨先烈和侨批家国记忆为三个清晰入口，通过问题式科普、人物故事和家书叙事降低理解门槛，同时坚持来源标注和待核实提示。",
    )

    add_heading(document, "7.3  模式创新", 2)
    add_body(
        document,
        "项目构建“社区课堂、场馆寻访、资料整理、网页传播”的实践链条。社区课堂提出问题，场馆寻访寻找答案，数字网页保存并继续传播成果，既符合团队真实实践顺序，也打通一次线下活动与长期公益服务之间的连接。",
    )

    add_heading(document, "7.4  技术创新", 2)
    add_body(
        document,
        "项目选择轻量、可维护的静态网页技术，以响应式设计适配手机和电脑，以数据文件驱动课程、场馆、文章、时间轴、资源和题库内容。技术服务于内容展示和后续维护，不堆砌复杂功能，突出计算机专业服务文化传播的实际价值。",
    )

    add_heading(document, "第八章  总结", 1)
    add_body(
        document,
        "“青声传薪，数扬侨魂”项目从社区课堂出发，在与青少年的真实交流中发现问题，又在场馆与史料中继续寻找答案，最终通过数字网页把一次次线下相遇整理为可以长期被看见、被使用的公共资源。项目的意义不在于功能数量，而在于能否让更多青少年看见跨越山海的家国记忆，并从华侨先辈和青年先烈的选择中理解责任与担当。",
    )
    add_body(
        document,
        "后续，团队将继续补充真实实践资料，核实历史内容来源，完善课件、照片、推文和学习资源入口，优化手机端浏览与知识互动，让平台成为学校、社区和青少年可以持续使用的华侨革命史公益科普窗口。",
    )


def set_update_fields(document):
    settings = document.settings._element
    node = settings.find(qn("w:updateFields"))
    if node is None:
        node = OxmlElement("w:updateFields")
        settings.append(node)
    node.set(qn("w:val"), "true")


def add_alt_text_to_logo(document):
    for inline in document.inline_shapes:
        doc_pr = inline._inline.docPr
        doc_pr.set("title", "华侨大学校徽与中英文校名")
        doc_pr.set("descr", "华侨大学校徽与中英文校名标识")


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(SOURCE, "r") as zf:
        logo_bytes = zf.read("word/media/image1.png")

    document = Document(SOURCE)
    remove_existing_body(document)
    configure_page(document)
    configure_styles(document)
    add_cover(document, logo_bytes)
    add_toc(document)
    add_content(document)
    set_update_fields(document)
    add_alt_text_to_logo(document)

    document.core_properties.title = "“青声传薪，数扬侨魂”华侨革命史数字科普实践平台项目申报书"
    document.core_properties.subject = "2026年暑期社会实践项目申报书"
    document.core_properties.author = "华侨大学计算机科学与技术学院实践团"
    document.core_properties.keywords = "华侨革命史,青少年科普,社会实践,数字平台"
    document.core_properties.comments = "按目录式项目申报书重新编排，不含预算、审批和签名表。"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
